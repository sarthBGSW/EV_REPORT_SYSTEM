import streamlit as st
import os
import sys
from dotenv import load_dotenv

# MUST be first Streamlit command
st.set_page_config(page_title="Azure AI - Agentic Document Generator", layout="wide")

# Load environment variables FIRST before any other imports
load_dotenv()

from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
import io

from modules.tools import process_uploaded_files
from workflow import app_graph

# Initialize session state for document persistence
if 'generated_document' not in st.session_state:
    st.session_state.generated_document = None
if 'generation_timestamp' not in st.session_state:
    st.session_state.generation_timestamp = None

st.title("🚙 EV Report Generator 2025 (Multi-Agent System)")
st.markdown("### Powered by GPT-5-Mini, Claude Sonnet 4.5, Grok-4 & Local Models")

# Show previously generated document if exists
if st.session_state.generated_document:
    st.info(f"📄 **Previous document available** - Generated at {st.session_state.generation_timestamp.strftime('%B %d, %Y at %H:%M')} ({len(st.session_state.generated_document):,} characters)")
    
    # Create DOCX for previous document
    doc_prev = Document()
    title_prev = doc_prev.add_heading('EV Report 2025', 0)
    title_prev.alignment = 1
    doc_prev.add_paragraph(f"Generated: {st.session_state.generation_timestamp.strftime('%B %d, %Y at %H:%M')}")
    doc_prev.add_paragraph("_" * 50)
    
    paragraphs_prev = st.session_state.generated_document.split('\n\n')
    for para in paragraphs_prev:
        if para.strip():
            if para.strip().startswith('#'):
                heading_text = para.strip().lstrip('#').strip()
                heading_level = min(len(para.strip()) - len(para.strip().lstrip('#')), 3)
                doc_prev.add_heading(heading_text, level=heading_level)
            else:
                doc_prev.add_paragraph(para.strip())
    
    bio_prev = io.BytesIO()
    doc_prev.save(bio_prev)
    bio_prev.seek(0)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button(
            "📥 Download Previous DOCX", 
            bio_prev.getvalue(), 
            file_name=f"EV_Report_{st.session_state.generation_timestamp.strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="prev_docx"
        )
    with col2:
        st.download_button(
            "📄 Download Previous MD", 
            st.session_state.generated_document, 
            file_name=f"EV_Report_{st.session_state.generation_timestamp.strftime('%Y%m%d_%H%M')}.md",
            key="prev_md"
        )
    with col3:
        if st.button("🗑️ Clear Previous Document"):
            st.session_state.generated_document = None
            st.session_state.generation_timestamp = None
            st.rerun()

# Sidebar for Config
with st.sidebar:
    st.header("Settings")
    uploaded_files = st.file_uploader("Upload Manual Reports (GPT-Ent/Gemini-Pro)", accept_multiple_files=True, type=['pdf', 'txt'])
    
    model_choice = st.selectbox("Primary Writer Model", ["gpt-5-mini", "claude-sonnet", "grok-4"])
    st.info("The system automatically cross-verifies using a different model than the writer.")

# Main Input
user_prompt = st.text_area("Enter Topic & Requirements", "Generate a comprehensive report on the Global EV Passenger Car Market, Trends, and Policies up to Dec 2025.", height=100)

if st.button("Start Agent Swarm"):
    try:
        with st.spinner("Initializing Agents... Reading Uploaded Docs..."):
            # 1. Process Manual Uploads
            context_text = ""
            if uploaded_files:
                try:
                    context_text = process_uploaded_files(uploaded_files)
                    st.success(f"Processed {len(uploaded_files)} manual documents.")
                except Exception as e:
                    st.error(f"Error processing files: {str(e)}")
                    raise
            
            # 2. Initialize State
            initial_state = {
                "topic": user_prompt,
                "uploaded_context": context_text,
                "outline": [],
                "current_chapter_index": 0,
                "current_chapter_content": "",
                "research_notes": "",
                "reviews": "",
                "final_document": ""
            }

            # 3. Run Graph
            progress_bar = st.progress(0)
            status_text = st.empty()
            final_output = st.empty()
            error_display = st.empty()
            
            # We stream the events to show progress
            current_step = 0
            final_state = initial_state.copy()  # Track the final state
            
            try:
                # Note: Depending on langgraph version, use .stream or .invoke
                # Increase recursion limit to handle multiple chapter iterations
                # 200 iterations = ~200 chapters which should be more than enough
                config = {"recursion_limit": 200}
                for output in app_graph.stream(initial_state, config=config):
                    for key, value in output.items():
                        # Update final_state with the latest output
                        final_state = value
                        
                        # Get chapter progress info
                        current_idx = value.get('current_chapter_index', 0)
                        total_chapters = len(value.get('outline', []))
                        
                        if key == "planner":
                            status_text.write(f"✅ Outline Generated: {len(value['outline'])} Chapters")
                            st.info(f"📋 Chapters planned: {', '.join(value['outline'][:5])}{'...' if len(value['outline']) > 5 else ''}")
                        elif key == "research":
                            status_text.write(f"🔍 Researching chapter {current_idx + 1}/{total_chapters}...")
                        elif key == "write":
                            chapter_title = value.get('outline', [''])[current_idx] if current_idx < total_chapters else 'Unknown'
                            status_text.write(f"✍️ Writing chapter {current_idx + 1}/{total_chapters}: {chapter_title}")
                        elif key == "review":
                            status_text.write(f"⚖️ Reviewing chapter {current_idx + 1}/{total_chapters}...")
                            
                            # Update Preview (show progress, not full content)
                            if 'final_document' in value:
                                doc_length = len(value['final_document'])
                                final_output.info(f"📝 Document length: {doc_length:,} characters")
                                
                    # Simple progress simulation
                    current_step += 5
                    if current_step > 100: current_step = 100
                    progress_bar.progress(current_step)

                # Extract final content from final_state (not initial_state!)
                final_content = final_state.get('final_document', '')
                
                # Store in session state for persistence
                st.session_state.generated_document = final_content
                st.session_state.generation_timestamp = datetime.now()
                
                # Generate DOCX file
                st.success(f"✅ Document Generation Complete! ({len(final_content):,} characters)")
                
                # Create DOCX document
                doc = Document()
                
                # Add title
                title = doc.add_heading('EV Report 2025', 0)
                title.alignment = 1  # Center alignment
                
                # Add metadata
                doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
                doc.add_paragraph(f"Topic: {user_prompt}")
                doc.add_paragraph("_" * 50)
                
                # Add content (split by paragraphs)
                # final_content already captured above from final_state
                
                # Split content and add to document
                paragraphs = final_content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        # Check if it's a heading (starts with #)
                        if para.strip().startswith('#'):
                            heading_text = para.strip().lstrip('#').strip()
                            heading_level = min(len(para.strip()) - len(para.strip().lstrip('#')), 3)
                            doc.add_heading(heading_text, level=heading_level)
                        else:
                            doc.add_paragraph(para.strip())
                
                # Save to BytesIO for download
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                
                # Download buttons for both formats
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        "📥 Download DOCX", 
                        bio.getvalue(), 
                        file_name=f"EV_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with col2:
                    st.download_button(
                        "📄 Download Markdown", 
                        final_content, 
                        file_name=f"EV_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.md"
                    )
                
                # Show preview option
                with st.expander("📖 Preview Generated Content"):
                    st.markdown(final_content)
            except Exception as graph_error:
                # Special handling for recursion limit - save what we have
                if "Recursion limit" in str(graph_error) or "GraphRecursionError" in str(graph_error):
                    error_display.warning(f"⚠️ Recursion limit reached. Saving document generated so far...")
                    print(f"[MAIN] Recursion limit hit. Attempting to save partial document.", file=sys.stderr)
                    
                    # Try to extract final_state if we have it
                    if 'final_state' in locals() and final_state.get('final_document'):
                        final_content = final_state['final_document']
                        st.session_state.generated_document = final_content
                        st.session_state.generation_timestamp = datetime.now()
                        
                        st.info(f"✅ Partial document saved: {len(final_content):,} characters")
                        st.warning("The document generation was stopped due to recursion limit. You can still download the content generated so far.")
                        
                        # Still show the DOCX generation section
                        # (the code below will handle it)
                    else:
                        st.error("Could not retrieve partial document. Try reducing the number of chapters or increasing recursion limit further.")
                else:
                    error_display.error(f"Error during workflow execution: {str(graph_error)}")
                    st.error("Check the terminal/console for detailed error logs.")
                raise
    except Exception as e:
        st.error(f"Critical error: {str(e)}")
        st.error("Please check the terminal output for detailed error information.")
        import traceback
        st.code(traceback.format_exc())